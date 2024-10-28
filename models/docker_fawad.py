from docx import Document

# Create a new Document
doc = Document()
doc.add_heading('Tutorial: Setting Up a Folder and Container, Moving Files, and Running a Script', 0)

# Introduction
doc.add_heading('1. Introduction', level=1)
doc.add_paragraph(
    """
    Objective:
    - Learn how to create and manage a folder on your host machine.
    - Transfer files and folders from the host to the Docker container.
    - Run Python scripts inside the Docker container.

    Pre-requisites:
    - Basic understanding of Docker and Python.
    - Docker installed on your machine.
    """
)

# Setting Up Your Host Environment
doc.add_heading('2. Setting Up Your Host Environment', level=1)

# Step 1: View All Folders on the Host
doc.add_heading('Step 1: View All Folders on the Host', level=2)
doc.add_paragraph(
    """
    Start by viewing all the folders in your home directory on the host machine:
    """)
doc.add_paragraph(
    """
    ls ~/
    """, style='Code'
)
doc.add_paragraph(
    "This command lists all the directories and files in your home directory."
)

# Step 2: Create a Folder Named `fawad_data`
doc.add_heading('Step 2: Create a Folder Named `fawad_data`', level=2)
doc.add_paragraph(
    """
    Create a new folder named `fawad_data` on your host machine:
    """)
doc.add_paragraph(
    """
    mkdir ~/fawad_data
    """, style='Code'
)
doc.add_paragraph(
    "This command creates a directory called `fawad_data` inside your home directory."
)

# Step 3: Confirm the Folder
doc.add_heading('Step 3: Confirm the Folder', level=2)
doc.add_paragraph(
    """
    Verify that the folder `fawad_data` has been created by listing the contents of your home directory again:
    """)
doc.add_paragraph(
    """
    ls ~/
    """, style='Code'
)
doc.add_paragraph(
    "You should see `fawad_data` listed among the other folders."
)

# Step 4: View All Docker Containers
doc.add_heading('Step 4: View All Docker Containers', level=2)
doc.add_paragraph(
    """
    To check the status of Docker containers on your system, run:
    """)
doc.add_paragraph(
    """
    docker ps -a
    """, style='Code'
)
doc.add_paragraph(
    "This command lists all Docker containers, showing both running and stopped containers."
)

# Step 4.1: View All Docker Images
doc.add_heading('Step 4.1: View All Docker Images', level=2)
doc.add_paragraph(
    """
    To see all Docker images available on your system, run:
    """)
doc.add_paragraph(
    """
    docker images
    """, style='Code'
)
doc.add_paragraph(
    "This command lists all Docker images that are available locally on your host machine. This is useful to ensure that the TensorFlow image you plan to use is available."
)

# Step 5: Create a Docker Container Named `fawad`
doc.add_heading('Step 5: Create a Docker Container Named `fawad`', level=2)
doc.add_paragraph(
    """
    Now, create a Docker container named `fawad` using the TensorFlow image, and make sure it uses GPU 2:
    """)
doc.add_paragraph(
    """
    docker run --gpus '"device=2"' -it --name fawad tensorflow/tensorflow:2.16.1-gpu bash
    """, style='Code'
)
doc.add_paragraph(
    """
    - `--gpus '"device=2"'`: Ensures that the container uses GPU 2.
    - `-it`: Opens an interactive terminal inside the container.
    - `--name fawad`: Names the container `fawad`.
    - `bash`: Starts a Bash shell in the container.
    """
)

# Step 6: Confirm the Container
doc.add_heading('Step 6: Confirm the Container', level=2)
doc.add_paragraph(
    """
    To confirm that the container `fawad` was created, exit the container and list all containers again:
    """)
doc.add_paragraph(
    """
    exit
    docker ps -a
    """, style='Code'
)
doc.add_paragraph(
    "You should see `fawad` listed among the containers."
)

# Step 7: Attach the Container to the TensorFlow Image
doc.add_heading('Step 7: Attach the Container to the TensorFlow Image', level=2)
doc.add_paragraph(
    """
    If you haven’t already attached to the TensorFlow image, you do so when creating the container. However, if you need to reattach to the container later, you can use:
    """)
doc.add_paragraph(
    """
    docker start -ai fawad
    """, style='Code'
)
doc.add_paragraph(
    "This command starts the container `fawad` and attaches you to the terminal session."
)

# Step 8: Create a File on Your Local Machine
doc.add_heading('Step 8: Create a File on Your Local Machine', level=2)
doc.add_paragraph(
    """
    **Step Explanation:** On your local machine (outside of the Docker container), create a Python script named `hello_world.py` inside the `fawad_data` directory.
    """
)

# Step 9: Move the File into WinSCP
doc.add_heading('Step 9: Move the File into WinSCP', level=2)
doc.add_paragraph(
    """
    **Step Explanation:** Use WinSCP (a file transfer tool) to move the `hello_world.py` file into the WinSCP window. This allows you to transfer files between your local machine and the host.
    """
)

# Step 10: Move the File to the Host Folder
doc.add_heading('Step 10: Move the File to the Host Folder', level=2)
doc.add_paragraph(
    """
    **Step Explanation:** Drag the `hello_world.py` file from WinSCP into the `fawad_data` folder on your host machine. This places the file on your server, ready to be copied into the Docker container.
    """
)

# Step 11: Copy the File to the Container
doc.add_heading('Step 11: Copy the File to the Container', level=2)
doc.add_paragraph(
    """
    Once the file is on your host machine, copy it into the Docker container using the following command:
    """)
doc.add_paragraph(
    """
    docker cp ~/fawad_data/hello_world.py fawad:/hello_world.py
    """, style='Code'
)
doc.add_paragraph(
    """
    - `~/fawad_data/hello_world.py`: The source file on your host machine.
    - `fawad:/hello_world.py`: The destination path inside the container.
    """
)

# Step 12: Run the Python Script Inside the Container
doc.add_heading('Step 12: Run the Python Script Inside the Container', level=2)
doc.add_paragraph(
    """
    Finally, run the Python script inside the `fawad` container:
    """)
doc.add_paragraph(
    """
    python hello_world.py
    """, style='Code'
)
doc.add_paragraph(
    """
    If you’re using Python 3, you might need to run:
    """)
doc.add_paragraph(
    """
    python3 hello_world.py
    """, style='Code'
)
doc.add_paragraph(
    "This will execute the script inside the container, and you should see the output in your terminal."
)

# Conclusion
doc.add_heading('Conclusion', level=1)
doc.add_paragraph(
    """
    This guide walks you through the entire process of creating a folder on your host machine, setting up a Docker container, transferring files using WinSCP, and running a Python script within the container. Each step ensures that you have a smooth and organized workflow between your local machine, the host, and the Docker container.
    """
)

# Save the document
doc_path = '/mnt/data/Docker_Tutorial_Fawad.docx'
doc.save(doc_path)
doc_path
